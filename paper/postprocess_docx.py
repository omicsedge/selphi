#!/usr/bin/env python3
"""Post-process the pandoc-built selphi2_paper.docx:
  1. remove ALL pandoc bookmarks; add one anchor per reference (ref1..refN);
  2. Times New Roman everywhere (docDefaults + every style + every run);
  3. body justified; figure images centered; figure captions centered + 10pt;
  4. tables: full width, grey borders, grey bold header + zebra grey rows,
     cells centred (H+V), header & data 10pt;
  5. in-text citation superscripts -> blue clickable hyperlinks to the matching
     reference (multi-cites like "7,8" link each number separately).
Usage: python3 postprocess_docx.py in.docx out.docx
"""
import re, sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

inp, outp = sys.argv[1], sys.argv[2]
doc = Document(inp)

FONT        = "Times New Roman"
HEADER_FILL = "595959"
ROW_LIGHT   = "FFFFFF"
ROW_DARK    = "E7E6E6"
BORDER      = "BFBFBF"
TABLE_SZ    = Pt(10)
LINK_BLUE   = "0563C1"

def W(tag): return qn('w:' + tag)
def set_fonts(rPr):
    rF = rPr.find(W('rFonts'))
    if rF is None: rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
    for a in ('ascii', 'hAnsi', 'cs', 'eastAsia'): rF.set(W(a), FONT)
def ptext(p): return ''.join(t.text or '' for t in p.iter(W('t')))

# --- 1a. remove all bookmarks ---
nbk = 0
for tag in ('bookmarkStart', 'bookmarkEnd'):
    for el in list(doc.element.iter(W(tag))):
        el.getparent().remove(el); nbk += 1

# --- 2. Times New Roman: docDefaults + every style ---
sroot = doc.styles.element
dd = sroot.find(W('docDefaults'))
if dd is None:
    dd = OxmlElement('w:docDefaults'); sroot.insert(0, dd)
rprd = dd.find(W('rPrDefault'))
if rprd is None: rprd = OxmlElement('w:rPrDefault'); dd.append(rprd)
rpr = rprd.find(W('rPr'))
if rpr is None: rpr = OxmlElement('w:rPr'); rprd.append(rpr)
set_fonts(rpr)
for st in sroot.findall(W('style')):
    srpr = st.find(W('rPr'))
    if srpr is None: srpr = OxmlElement('w:rPr'); st.append(srpr)
    set_fonts(srpr)

base = (doc.styles["Normal"].font.size or Pt(12)) if "Normal" in [s.name for s in doc.styles] else Pt(12)
cap_sz = Pt(max(9, int(base.pt) - 2))

def has_image(p): return len(p._p.findall('.//' + W('drawing'))) > 0
def is_caption(p): return p.text.strip().startswith(("Figure 1.", "Figure 2.", "Figure 3."))
def set_run(r, size=None, bold=None, white=False):
    set_fonts(r._element.get_or_add_rPr())
    if size is not None: r.font.size = size
    if bold is not None: r.font.bold = bold
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# --- 3. body paragraphs: alignment + fonts ---
nimg = njust = ncap = 0
for p in doc.paragraphs:
    if has_image(p):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; nimg += 1
        for r in p.runs: set_run(r)
    elif is_caption(p):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; ncap += 1
        for r in p.runs: set_run(r, size=cap_sz)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; njust += 1
        for r in p.runs: set_run(r)

# --- 4. tables ---
def set_cell_bg(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(W('shd')): tcPr.remove(e)
    shd = OxmlElement('w:shd'); shd.set(W('val'), 'clear'); shd.set(W('color'), 'auto'); shd.set(W('fill'), hexfill)
    tcPr.append(shd)
def full_width_and_borders(t):
    tblPr = t._tbl.tblPr
    for e in tblPr.findall(W('tblW')): tblPr.remove(e)
    tw = OxmlElement('w:tblW'); tw.set(W('type'), 'pct'); tw.set(W('w'), '5000'); tblPr.append(tw)
    for e in tblPr.findall(W('tblBorders')): tblPr.remove(e)
    bd = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(W('val'), 'single'); b.set(W('sz'), '4'); b.set(W('space'), '0'); b.set(W('color'), BORDER)
        bd.append(b)
    tblPr.append(bd)
# usable text width (twips) = page width - margins; fallback US-Letter 1in margins
sec = doc.sections[0]
try:
    usable_tw = int((sec.page_width - sec.left_margin - sec.right_margin) / 635)
except Exception:
    usable_tw = 9360
if not usable_tw or usable_tw < 3000:
    usable_tw = 9360
FLOOR = 720  # min column width (~0.5 in) so short numeric/label cells don't wrap

CHARW = 105  # ~twips per char at 10pt Times; PAD = cell L/R margins + slack
def col_metrics(t):
    ncol = len(t.rows[0].cells)
    data_ml = [0] * ncol     # longest DATA cell (drives width; headers may wrap)
    word_ml = [0] * ncol     # longest single word (so headers wrap at spaces, not mid-word)
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            if ci >= ncol: continue
            txt = cell.text.strip()
            if ri > 0:
                data_ml[ci] = max(data_ml[ci], len(txt))
            for wd in txt.split():
                word_ml[ci] = max(word_ml[ci], len(wd))
    nat = [max(data_ml[i], word_ml[i]) * CHARW + 260 for i in range(ncol)]  # no-wrap on data
    wide = [data_ml[i] > 20 for i in range(ncol)]   # text columns (by DATA) get the remainder
    if any(wide):
        narrow_sum = sum(nat[i] for i in range(ncol) if not wide[i])
        rem = max(usable_tw - narrow_sum, 1600 * sum(wide))
        wtot = sum(data_ml[i] for i in range(ncol) if wide[i]) or 1
        cw = [ (int(rem * data_ml[i] / wtot) if wide[i] else nat[i]) for i in range(ncol) ]
    else:                                            # all short: scale naturals to fill the page
        tot = sum(nat) or 1
        cw = [int(usable_tw * n / tot) for n in nat]
    cw = [max(c, FLOOR) for c in cw]
    cw[data_ml.index(max(data_ml))] += usable_tw - sum(cw)   # widest-data column absorbs slack
    return data_ml, cw

ntab = 0
for t in doc.tables:
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    full_width_and_borders(t)
    ml, cw = col_metrics(t); ncol = len(cw)
    tblPr = t._tbl.tblPr
    for e in tblPr.findall(W('tblLayout')): tblPr.remove(e)
    tl = OxmlElement('w:tblLayout'); tl.set(W('type'), 'fixed'); tblPr.append(tl)
    for e in tblPr.findall(W('tblW')): tblPr.remove(e)
    tw = OxmlElement('w:tblW'); tw.set(W('type'), 'dxa'); tw.set(W('w'), str(sum(cw))); tblPr.append(tw)
    grid = t._tbl.find(W('tblGrid'))
    if grid is not None:
        for i, gc in enumerate(grid.findall(W('gridCol'))):
            if i < ncol: gc.set(W('w'), str(cw[i]))
    left_col = [m > 18 for m in ml]                # text-heavy columns -> left aligned
    for ri, row in enumerate(t.rows):
        header = (ri == 0)
        fill = HEADER_FILL if header else (ROW_DARK if ri % 2 == 0 else ROW_LIGHT)
        for ci, cell in enumerate(row.cells):
            set_cell_bg(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcPr = cell._tc.get_or_add_tcPr()
            for e in tcPr.findall(W('tcW')): tcPr.remove(e)
            cwe = OxmlElement('w:tcW'); cwe.set(W('type'), 'dxa'); cwe.set(W('w'), str(cw[ci] if ci < ncol else FLOOR)); tcPr.append(cwe)
            align = WD_ALIGN_PARAGRAPH.LEFT if (not header and ci < ncol and left_col[ci]) else WD_ALIGN_PARAGRAPH.CENTER
            for p in cell.paragraphs:
                p.alignment = align
                for r in p.runs: set_run(r, size=TABLE_SZ, bold=True if header else None, white=header)
    ntab += 1

# --- 5. citation superscripts -> blue clickable links to the references ---
def sup_link_run(text, anchor=None):
    """A superscript blue run (wrapped in a hyperlink if anchor given)."""
    r = OxmlElement('w:r'); rp = OxmlElement('w:rPr')
    set_fonts(rp)
    c = OxmlElement('w:color'); c.set(W('val'), LINK_BLUE); rp.append(c)
    va = OxmlElement('w:vertAlign'); va.set(W('val'), 'superscript'); rp.append(va)
    r.append(rp)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    if anchor:
        h = OxmlElement('w:hyperlink'); h.set(W('anchor'), anchor); h.append(r); return h
    return r

body = doc.element.body
paras = [c for c in body if c.tag == W('p')]
def find_idx(name):
    for i, p in enumerate(paras):
        if ptext(p).strip() == name: return i
    return None
abs_i, ref_i = find_idx('Abstract'), find_idx('References')

# 1b. anchor each reference (numbered list items after the References heading,
#     stopping at the next heading so appended Supplementary content is untouched)
def pstyle(p):
    pPr = p.find(W('pPr'))
    st = pPr.find(W('pStyle')) if pPr is not None else None
    return st.get(W('val')) if st is not None else ''
nref = 0
if ref_i is not None:
    for p in paras[ref_i + 1:]:
        if nref > 0 and (pstyle(p).startswith('Heading') or pstyle(p) == 'Title'):
            break
        if p.find('.//' + W('numPr')) is not None:
            nref += 1
            bs = OxmlElement('w:bookmarkStart'); bs.set(W('id'), str(4000 + nref)); bs.set(W('name'), f'ref{nref}')
            be = OxmlElement('w:bookmarkEnd'); be.set(W('id'), str(4000 + nref))
            pPr = p.find(W('pPr'))
            (pPr.addnext(bs) if pPr is not None else p.insert(0, bs)); p.append(be)

# link citations in body paragraphs (between Abstract and References)
nlink = 0
lo = (abs_i or 0); hi = (ref_i if ref_i is not None else len(paras))
CIT = re.compile(r'^\d[\d,]*$')
for p in paras[lo:hi]:
    for r in list(p):
        if r.tag != W('r'): continue
        rpr = r.find(W('rPr'))
        if rpr is None or rpr.find(W('vertAlign')) is None: continue
        if rpr.find(W('vertAlign')).get(W('val')) != 'superscript': continue
        te = r.find(W('t'));  txt = (te.text if te is not None else '') or ''
        if not CIT.match(txt.strip()): continue
        idx = list(p).index(r); p.remove(r)
        for j, tok in enumerate(txt.strip().split(',')):
            if j: p.insert(idx, sup_link_run(',')); idx += 1
            p.insert(idx, sup_link_run(tok, anchor=f'ref{tok}' if tok.isdigit() else None)); idx += 1
            nlink += 1

# page break before the Supplementary Information section (if present)
npb = 0
for p in paras:
    if ptext(p).strip() == 'Supplementary Information':
        pPr = p.find(W('pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr'); p.insert(0, pPr)
        if pPr.find(W('pageBreakBefore')) is None:
            pPr.insert(0, OxmlElement('w:pageBreakBefore')); npb = 1
        break

doc.save(outp)
print(f"bookmarks removed: {nbk}; font->{FONT}; images centered: {nimg}; justified: {njust}; "
      f"captions {cap_sz.pt:.0f}pt: {ncap}; tables: {ntab}; ref anchors: {nref}; citation links: {nlink}; supp page-break: {npb}")
